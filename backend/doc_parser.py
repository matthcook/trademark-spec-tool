import re
from docx import Document
from dataclasses import dataclass
from typing import Optional


@dataclass
class FormattedRun:
    text: str
    bold: bool
    underline: bool


@dataclass
class ParsedOfficeAction:
    application_number: Optional[str]
    ir_number: Optional[str]          # International Registration number
    trademark_name: Optional[str]
    applicant_name: Optional[str]
    formatting_convention: Optional[str]
    full_text: str
    paragraphs_with_formatting: list[dict]


def _iter_header_chunks(doc) -> list[str]:
    """Extract text from all page header sections (not in doc.paragraphs/doc.tables)."""
    chunks = []
    try:
        for section in doc.sections:
            for hdr in [section.header, section.first_page_header]:
                if hdr is None:
                    continue
                for para in hdr.paragraphs:
                    if para.text.strip():
                        chunks.append(para.text.strip())
                for tbl in hdr.tables:
                    for row in tbl.rows:
                        cells = [c.text.strip() for c in row.cells]
                        chunks.extend(c for c in cells if c)
                        for i in range(len(cells) - 1):
                            if cells[i] and cells[i + 1]:
                                chunks.append(f"{cells[i]}: {cells[i + 1]}")
    except Exception:
        pass
    return chunks


def _iter_textbox_chunks(doc) -> list[str]:
    """Extract text from text boxes (w:txbxContent elements), which are invisible to doc.paragraphs."""
    chunks = []
    try:
        W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        for txbx in doc.element.body.iter(f"{{{W}}}txbxContent"):
            for p in txbx.iter(f"{{{W}}}p"):
                text = "".join(t.text or "" for t in p.iter(f"{{{W}}}t")).strip()
                if text:
                    chunks.append(text)
    except Exception:
        pass
    return chunks


def _all_text_chunks(paragraphs: list[dict], doc) -> list[str]:
    """Collect every visible text chunk: body paragraphs, body tables, headers, text boxes."""
    chunks = [p["text"] for p in paragraphs if p["text"].strip()]

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            chunks.extend(c for c in cells if c)
            for i in range(len(cells) - 1):
                if cells[i] and cells[i + 1]:
                    chunks.append(f"{cells[i]}: {cells[i + 1]}")

    chunks.extend(_iter_header_chunks(doc))
    chunks.extend(_iter_textbox_chunks(doc))
    return chunks


def _clean_num(s: str) -> Optional[str]:
    n = re.sub(r"[\s,\xa0]", "", s)   # also strip non-breaking spaces
    return n if re.fullmatch(r"\d{7,8}", n) else None


def _find_formatted_numbers(text: str) -> list[str]:
    """
    Find all 7-8 digit numbers, including ones formatted with spaces/commas
    as thousand separators (e.g. "2 236 993" or "2,236,993").
    """
    results = []
    for m in re.finditer(r"(?<!\d)([\d][\d,\s\xa0]{4,9}[\d])(?!\d)", text):
        cleaned = _clean_num(m.group(1))
        if cleaned:
            results.append(cleaned)
    return results


def _extract_numbers(paragraphs: list[dict], doc) -> tuple[Optional[str], Optional[str]]:
    """
    Return (cipo_app_number, ir_number).

    Searches body paragraphs, body tables, page headers, and text boxes.

    CIPO app number heuristics (in priority order):
      1. Value after a label: "Our File", "Our Ref", "Notre dossier", "File No.", etc.
      2. Any 7-digit number starting with 2 (CIPO app numbers ~2,000,000–2,200,000)
         — handles plain "2236993" and space-formatted "2 236 993"
      3. Any 7-8 digit number as last resort

    IR number: value after "IR No", "IR Number", "International Registration", etc.
    """
    chunks = _all_text_chunks(paragraphs, doc)

    our_file_re = re.compile(
        r"(?:"
        r"Our\s+(?:File|Ref(?:erence)?)"
        r"|Notre\s+(?:dossier|r[ée]f(?:[ée]rence)?)"
        r"|No\.?\s*de\s+(?:dossier|r[ée]f(?:[ée]rence)?)"
        r"|File\s+No\.?"
        r"|Ref(?:erence)?\s*No\.?"
        r")"
        r"[^0-9]{0,30}?([\d][\d,\s\xa0]{4,9}[\d])",
        re.IGNORECASE,
    )
    ir_re = re.compile(
        r"(?:IR\s+(?:No\.?|Number|Num\.?)|"
        r"Int(?:ernational)?\s+Reg(?:istration)?\.?\s*(?:No\.?|Number)?)"
        r"[^0-9]{0,20}?([\d][\d,\s\xa0]{4,9}[\d])",
        re.IGNORECASE,
    )

    app_number: Optional[str] = None
    ir_number: Optional[str] = None

    # Pass 1 — labeled searches
    for text in chunks:
        if not app_number:
            m = our_file_re.search(text)
            if m:
                app_number = _clean_num(m.group(1))
        if not ir_number:
            m = ir_re.search(text)
            if m:
                ir_number = _clean_num(m.group(1))
        if app_number and ir_number:
            break

    # Pass 2 — any 7-digit number starting with 2 (handles spaced format too)
    if not app_number:
        for text in chunks:
            for candidate in _find_formatted_numbers(text):
                if re.fullmatch(r"2\d{6}", candidate) and candidate != ir_number:
                    app_number = candidate
                    break
            if app_number:
                break

    # Pass 3 — any 7-8 digit number
    if not app_number:
        for text in chunks:
            for candidate in _find_formatted_numbers(text):
                if candidate != ir_number:
                    app_number = candidate
                    break
            if app_number:
                break

    return app_number, ir_number


def _extract_re_table(doc) -> dict:
    """Extract trademark name and applicant from the RE: header table or RE: paragraph."""
    result = {"trademark_name": None, "applicant_name": None}

    # Try table-based extraction first
    for table in doc.tables:
        cells = [cell.text.strip() for row in table.rows for cell in row.cells if cell.text.strip()]
        if any(c in ("RE:", "Re:", "Trademark:", "Applicant:") for c in cells):
            for i, cell in enumerate(cells):
                if cell in ("Trademark:", "Marque de commerce:") and i + 1 < len(cells):
                    result["trademark_name"] = cells[i + 1].replace("\n", " ").strip()
                if cell in ("Applicant:", "Demandeur:") and i + 1 < len(cells):
                    result["applicant_name"] = cells[i + 1].replace("\n", " ").strip()

    # Fallback: tab-separated RE: paragraphs e.g. "RE:\tTrademark:\tWeather"
    if not result["trademark_name"] or not result["applicant_name"]:
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            parts = [p.strip() for p in re.split(r"\t+", text)]
            # "RE:\tTrademark:\tWeather" or "Trademark:\tWeather"
            for i, part in enumerate(parts):
                if not result["trademark_name"] and re.match(r"Trademark\s*:|Marque\s*de\s*commerce\s*:", part, re.I):
                    if i + 1 < len(parts) and parts[i + 1]:
                        result["trademark_name"] = parts[i + 1]
                if not result["applicant_name"] and re.match(r"Applicant\s*:|Demandeur\s*:", part, re.I):
                    if i + 1 < len(parts) and parts[i + 1]:
                        result["applicant_name"] = parts[i + 1]

    return result


def extract_document_debug(file_path: str) -> dict:
    """Return raw text from all document areas for diagnosing extraction issues."""
    doc = Document(file_path)
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    tables = []
    for t_idx, table in enumerate(doc.tables):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        tables.append({"table_index": t_idx, "rows": rows})
    headers = _iter_header_chunks(doc)
    textboxes = _iter_textbox_chunks(doc)
    return {
        "paragraphs": paras[:60],
        "tables": tables,
        "headers": headers,
        "textboxes": textboxes,
    }


def parse_office_action(file_path: str) -> ParsedOfficeAction:
    doc = Document(file_path)
    paragraphs_with_formatting = []
    full_text_lines = []

    for para in doc.paragraphs:
        if not para.text.strip():
            continue

        tagged = ""
        runs = []
        for run in para.runs:
            t = run.text
            bold = bool(run.bold)
            underline = bool(run.underline)
            if bold and underline:
                t = f"[BOLD_UNDERLINE]{t}[/BOLD_UNDERLINE]"
            elif bold:
                t = f"[BOLD]{t}[/BOLD]"
            elif underline:
                t = f"[UNDERLINE]{t}[/UNDERLINE]"
            tagged += t
            runs.append({"text": run.text, "bold": bold, "underline": underline})

        paragraphs_with_formatting.append({
            "text": para.text,
            "tagged": tagged,
            "runs": runs,
        })
        full_text_lines.append(para.text)

    re_fields = _extract_re_table(doc)
    app_number, ir_number = _extract_numbers(paragraphs_with_formatting, doc)

    return ParsedOfficeAction(
        application_number=app_number,
        ir_number=ir_number,
        trademark_name=re_fields["trademark_name"],
        applicant_name=re_fields["applicant_name"],
        formatting_convention=None,
        full_text="\n".join(full_text_lines),
        paragraphs_with_formatting=paragraphs_with_formatting,
    )
